"""
Job Scraper Pro  v4
═══════════════════════════════════════════════════════
Sites   : Indeed · LinkedIn (API) · Glassdoor · InfoJobs · Subito · Monster IT
Export  : CSV + Word (.docx)
Notify  : Telegram Bot · Email (SSL/TLS auto)
GUI     : Modern dark sidebar + light content area
Fixes   : LinkedIn via JSON API · DOCX float bug · Email SSL 465
═══════════════════════════════════════════════════════
pip install selenium pandas python-docx requests
"""

import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox
import threading, os, re, time, smtplib, json, ssl
from datetime import datetime
from urllib.parse import urlencode, quote_plus
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

import pandas as pd
import requests

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ── files — always in same folder as this script ─────────────────────────────
_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CSV_FILE   = os.path.join(_BASE_DIR, "jobs.csv")
DOCX_FILE  = os.path.join(_BASE_DIR, "jobs_report.docx")
XLSX_FILE  = os.path.join(_BASE_DIR, "jobs_report.xlsx")
CONFIG_FILE= os.path.join(_BASE_DIR, "scraper_config.json")

INDEED_DOMAINS = {
    "italy":"it.indeed.com","italia":"it.indeed.com",
    "uk":"uk.indeed.com","united kingdom":"uk.indeed.com",
    "germany":"de.indeed.com","france":"fr.indeed.com",
    "spain":"es.indeed.com","canada":"ca.indeed.com",
    "australia":"au.indeed.com","usa":"www.indeed.com",
    "united states":"www.indeed.com",
}

# ── design tokens ─────────────────────────────────────────────────────────────
SIDEBAR   = "#1A1F2E"
SIDEBAR2  = "#242938"
ACCENT    = "#4F8EF7"
ACCENT2   = "#3B7AE8"
SUCCESS   = "#34D399"
WARNING   = "#FBBF24"
DANGER    = "#F87171"
BG        = "#F0F4F8"
CARD      = "#FFFFFF"
BORDER    = "#E2E8F0"
TXT       = "#1E293B"
TXT2      = "#475569"
MUTED     = "#94A3B8"
LOG_BG    = "#0F172A"
LOG_TXT   = "#94A3B8"


# ══════════════════════════════════════════════════════════════════════════════
#  CONFIG
# ══════════════════════════════════════════════════════════════════════════════

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            return json.load(open(CONFIG_FILE, encoding="utf-8"))
        except Exception:
            pass
    return {}

def save_config(cfg):
    json.dump(cfg, open(CONFIG_FILE, "w", encoding="utf-8"), indent=2, ensure_ascii=False)


# ══════════════════════════════════════════════════════════════════════════════
#  CSV
# ══════════════════════════════════════════════════════════════════════════════

def load_existing_links():
    if os.path.exists(CSV_FILE):
        df = pd.read_csv(CSV_FILE)
        if "Link" in df.columns:
            return set(df["Link"].dropna().astype(str))
    return set()

def save_to_csv(jobs, run_id=None):
    """
    Save jobs to CSV.
    Same link CAN appear in multiple runs (user wants to see it each time).
    We only deduplicate within the SAME run_id.
    Returns count of rows added.
    """
    if not jobs:
        return 0

    # Tag every job with this run's timestamp
    if run_id is None:
        run_id = datetime.now().strftime("%Y-%m-%d %H:%M")
    for j in jobs:
        j["Run"] = run_id

    new_df = pd.DataFrame(jobs)

    if os.path.exists(CSV_FILE):
        old_df = pd.read_csv(CSV_FILE)
        # Only deduplicate within this run (same Link + same Run)
        new_df = new_df.drop_duplicates(subset=["Link"])
        combined = pd.concat([old_df, new_df], ignore_index=True)
        combined.to_csv(CSV_FILE, index=False)
        return len(new_df)
    else:
        new_df = new_df.drop_duplicates(subset=["Link"])
        new_df.to_csv(CSV_FILE, index=False)
        return len(new_df)


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX  — float bug fixed: str(val or "")
# ══════════════════════════════════════════════════════════════════════════════

def _safe(val):
    """Convert any value (including NaN/float) to clean string."""
    if val is None:
        return ""
    try:
        import math
        if isinstance(val, float) and math.isnan(val):
            return ""
    except Exception:
        pass
    return str(val).strip()

def _shd(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

def save_to_xlsx(jobs, log):
    """Save jobs list to Excel with formatting."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        log("[Excel] openpyxl not installed — run: pip install openpyxl")
        return ""

    if not jobs:
        log("[Excel] no jobs to save")
        return ""

    xlsx_path = XLSX_FILE
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Jobs"

    headers = ["#", "Title", "Company", "Location", "Source", "Date Found", "Link"]
    col_widths = [5, 40, 25, 20, 12, 18, 60]

    # Header style
    hdr_fill = PatternFill("solid", fgColor="2563EB")
    hdr_font = Font(bold=True, color="FFFFFF", size=10)
    thin = Side(style="thin", color="D1D9E6")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for ci, (h, w) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=1, column=ci, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border
        ws.column_dimensions[get_column_letter(ci)].width = w

    ws.row_dimensions[1].height = 20

    # Data rows
    even_fill = PatternFill("solid", fgColor="EFF6FF")
    link_font = Font(color="2563EB", underline="single", size=9)
    normal_font = Font(size=9)

    for ri, job in enumerate(jobs, 2):
        fill = even_fill if ri % 2 == 0 else PatternFill()
        vals = [
            ri - 1,
            _safe(job.get("Title")),
            _safe(job.get("Company")),
            _safe(job.get("Location")),
            _safe(job.get("Source")),
            _safe(job.get("Date Found")),
            _safe(job.get("Link")),
        ]
        for ci, val in enumerate(vals, 1):
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.border = border
            cell.fill = fill
            cell.alignment = Alignment(vertical="center", wrap_text=(ci == 2))
            # Link column: hyperlink
            if ci == 7 and val.startswith("http"):
                cell.hyperlink = val
                cell.font = link_font
            elif ci == 2:
                cell.font = Font(bold=True, size=9)
            else:
                cell.font = normal_font

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:G{len(jobs)+1}"

    wb.save(xlsx_path)
    log(f"[Excel] ✓ Saved → {os.path.abspath(xlsx_path)}")
    return os.path.abspath(xlsx_path)


def save_to_docx(jobs, log):
    if not DOCX_OK:
        log("[DOCX] python-docx not installed")
        return ""
    if not jobs:
        log("[DOCX] no jobs to save")
        return ""

    doc = DocxDocument()

    # Page margins
    from docx.shared import Cm
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    # Title
    h = doc.add_heading("", 0)
    h.clear()
    run = h.add_run("Job Scraper Report")
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    run.font.size = Pt(20)
    run.font.bold = True

    p = doc.add_paragraph()
    r = p.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  Total: {len(jobs)} jobs")
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    r.font.size = Pt(9)
    doc.add_paragraph()

    # Table
    headers = ["#", "Title", "Company", "Location", "Source", "Date"]
    col_w   = [Inches(0.28), Inches(2.6), Inches(1.4), Inches(1.2), Inches(0.75), Inches(0.85)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for i, h_txt in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h_txt
        _shd(cell, "1E3A8A")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for idx, job in enumerate(jobs, 1):
        row = table.add_row()
        vals = [
            str(idx),
            _safe(job.get("Title")),
            _safe(job.get("Company")),
            _safe(job.get("Location")),
            _safe(job.get("Source")),
            _safe(job.get("Date Found", ""))[:10],
        ]
        fill = "EFF6FF" if idx % 2 == 0 else "FFFFFF"

        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            _shd(cell, fill)

            # Title cell: add hyperlink
            if ci == 1:
                link = _safe(job.get("Link"))
                cell.paragraphs[0].clear()
                if link:
                    try:
                        rel = doc.part.relate_to(
                            link,
                            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                            is_external=True
                        )
                        hl = OxmlElement("w:hyperlink")
                        hl.set(qn("r:id"), rel)
                        wr = OxmlElement("w:r")
                        rPr = OxmlElement("w:rPr")
                        col_el = OxmlElement("w:color"); col_el.set(qn("w:val"), "2563EB")
                        u_el   = OxmlElement("w:u");   u_el.set(qn("w:val"), "single")
                        sz_el  = OxmlElement("w:sz");  sz_el.set(qn("w:val"), "17")
                        rPr.append(col_el); rPr.append(u_el); rPr.append(sz_el)
                        wr.append(rPr)
                        t_el = OxmlElement("w:t"); t_el.text = val
                        wr.append(t_el)
                        hl.append(wr)
                        cell.paragraphs[0]._p.append(hl)
                    except Exception:
                        cell.paragraphs[0].add_run(val).font.size = Pt(8.5)
                else:
                    cell.paragraphs[0].add_run(val).font.size = Pt(8.5)
            else:
                r = cell.paragraphs[0].add_run(val)
                r.font.size = Pt(8.5)
                if ci == 0:
                    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Column widths
    for ci, w in enumerate(col_w):
        for row in table.rows:
            row.cells[ci].width = w

    doc.save(DOCX_FILE)
    log(f"[DOCX] ✓ Saved → {os.path.abspath(DOCX_FILE)}")
    return os.path.abspath(DOCX_FILE)


# ══════════════════════════════════════════════════════════════════════════════
#  EMAIL
# ══════════════════════════════════════════════════════════════════════════════

def send_email(jobs, cfg, log):
    try:
        host   = cfg.get("smtp_host", "smtp.gmail.com")
        port   = int(cfg.get("smtp_port", 465))
        user   = cfg.get("smtp_user", "")
        passwd = cfg.get("smtp_pass", "")
        to     = cfg.get("email_to", "") or user

        if not user or not passwd:
            log("[Email] ⚠ No credentials set")
            return False

        lines = [f"Job Scraper found {len(jobs)} job(s):\n"]
        for j in jobs:
            lines.append(f"• {_safe(j.get('Title'))} @ {_safe(j.get('Company'))} ({_safe(j.get('Location'))}) [{_safe(j.get('Source'))}]")
            lines.append(f"  {_safe(j.get('Link'))}\n")

        msg = MIMEMultipart()
        msg["From"] = user; msg["To"] = to
        msg["Subject"] = f"[Job Scraper] {len(jobs)} jobs found"
        msg.attach(MIMEText("\n".join(lines), "plain", "utf-8"))

        log(f"[Email] Connecting {host}:{port} …")
        ctx = ssl.create_default_context()
        if port == 465:
            with smtplib.SMTP_SSL(host, port, context=ctx, timeout=20) as s:
                s.login(user, passwd)
                s.sendmail(user, to, msg.as_string())
        else:
            with smtplib.SMTP(host, port, timeout=20) as s:
                s.ehlo(); s.starttls(context=ctx); s.ehlo()
                s.login(user, passwd)
                s.sendmail(user, to, msg.as_string())

        log(f"[Email] ✓ Sent to {to}")
        return True
    except smtplib.SMTPAuthenticationError:
        log("[Email] ✗ Auth failed — use App Password not your Gmail password")
        log("[Email]   myaccount.google.com → Security → App Passwords")
        return False
    except Exception as e:
        log(f"[Email] ✗ {e}")
        return False


# ══════════════════════════════════════════════════════════════════════════════
#  TELEGRAM
# ══════════════════════════════════════════════════════════════════════════════

def send_telegram(jobs, cfg, log):
    try:
        token   = cfg.get("tg_token", "").strip()
        chat_id = cfg.get("tg_chat_id", "").strip()
        if not token or not chat_id:
            log("[Telegram] ⚠ Token/ChatID missing")
            return False

        chunks = []
        current = f"🔍 *Job Scraper* — {len(jobs)} job(s) found\n\n"
        for j in jobs:
            line = (f"📌 *{_safe(j.get('Title'))}*\n"
                    f"🏢 {_safe(j.get('Company'))}  📍 {_safe(j.get('Location'))}  🌐 {_safe(j.get('Source'))}\n"
                    f"🔗 {_safe(j.get('Link'))}\n\n")
            if len(current) + len(line) > 3800:
                chunks.append(current)
                current = line
            else:
                current += line
        chunks.append(current)

        url = f"https://api.telegram.org/bot{token}/sendMessage"
        for i, chunk in enumerate(chunks):
            r = requests.post(url, json={
                "chat_id": chat_id, "text": chunk,
                "parse_mode": "Markdown", "disable_web_page_preview": True
            }, timeout=15)
            if not r.ok:
                log(f"[Telegram] ✗ chunk {i+1}: {r.text[:120]}")
                return False

        log(f"[Telegram] ✓ Sent {len(chunks)} message(s)")
        return True
    except Exception as e:
        log(f"[Telegram] ✗ {e}")
        return False


# ── IT job title filter ───────────────────────────────────────────────────────
# کلمات مرتبط با IT که در عنوان شغل باید باشن
def title_matches_keyword(title: str, keyword: str) -> bool:
    """فقط چک میکنه کلمات keyword توی عنوان هستن یا نه."""
    if not keyword or not title:
        return True
    t = title.lower()
    kw = keyword.lower().strip()
    # Exact phrase
    if kw in t:
        return True
    # All words present
    words = [w for w in kw.split() if len(w) > 1]
    if words and all(w in t for w in words):
        return True
    return False




def make_driver(headless=True):
    opts = webdriver.ChromeOptions()
    if headless:
        opts.add_argument("--headless=new")
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    opts.add_argument("--disable-blink-features=AutomationControlled")
    opts.add_experimental_option("excludeSwitches", ["enable-automation"])
    opts.add_experimental_option("useAutomationExtension", False)
    opts.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                      "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36")
    drv = webdriver.Chrome(options=opts)
    drv.execute_cdp_cmd("Page.addScriptToEvaluateOnNewDocument",
        {"source": "Object.defineProperty(navigator,'webdriver',{get:()=>undefined})"})
    return drv

def scroll_down(drv, times=3, pause=1.5):
    for _ in range(times):
        drv.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(pause)

def now_str():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def safe_str(val):
    return _safe(val)


# ══════════════════════════════════════════════════════════════════════════════
#  SCRAPERS
# ══════════════════════════════════════════════════════════════════════════════

def scrape_indeed(driver, keywords, country, city, existing, log, date_range=14, max_pages=10):
    jobs = []
    domain   = INDEED_DOMAINS.get(country.strip().lower(), "it.indeed.com")
    # Build broader query — use each keyword separately joined with space
    query    = " ".join(keywords)
    location = city if city.strip() else country
    base     = {"q": query, "l": location, "sort": "date"}
    if date_range > 0:
        base["fromage"] = str(date_range)

    # Track only links found THIS run to avoid within-run dupes
    # (don't skip links already in CSV — user may want to see them again)
    seen_this_run = set()

    for page in range(max_pages):
        params = dict(base)
        if page > 0:
            params["start"] = str(page * 10)
        url = f"https://{domain}/jobs?{urlencode(params)}"
        log(f"[Indeed] p{page+1}: {url}")
        try:
            driver.get(url)
            try:
                WebDriverWait(driver, 15).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR,
                        "a[id^='job_'], .jobsearch-NoResult, #mosaic-provider-jobcards")))
            except Exception:
                pass
            time.sleep(2)

            if driver.find_elements(By.CSS_SELECTOR, ".jobsearch-NoResult"):
                log(f"[Indeed] no results on p{page+1}")
                break

            # Try multiple selectors — Indeed changes structure frequently
            anchors = driver.find_elements(By.CSS_SELECTOR, "a[id^='job_']")
            if not anchors:
                anchors = driver.find_elements(By.CSS_SELECTOR, "a[data-jk]")
            if not anchors:
                # Extract job links from any anchor with jk= in href
                all_links = driver.find_elements(By.CSS_SELECTOR, "a[href*='jk=']")
                anchors = [a for a in all_links if a.get_attribute("href")]
            if not anchors:
                # Last resort: get job cards and find links inside
                cards = driver.find_elements(By.CSS_SELECTOR,
                    "div[data-jk], li[data-jk], div.job_seen_beacon")
                for card in cards:
                    a_els = card.find_elements(By.CSS_SELECTOR, "a[href]")
                    anchors.extend(a_els)

            log(f"[Indeed] p{page+1}: {len(anchors)} anchors")
            if not anchors:
                # Log page source snippet for debugging
                try:
                    body = driver.find_element(By.TAG_NAME, "body").text[:200]
                    log(f"[Indeed] page preview: {body[:100]}")
                except Exception:
                    pass
                break

            page_new = 0
            for a in anchors:
                try:
                    jk = ""
                    # Method 1: id="job_XXXX"
                    aid = a.get_attribute("id") or ""
                    if aid.startswith("job_"):
                        jk = aid.replace("job_", "").strip()
                    # Method 2: data-jk attribute
                    if not jk:
                        jk = a.get_attribute("data-jk") or ""
                    # Method 3: jk= in href
                    if not jk:
                        href = a.get_attribute("href") or ""
                        m = re.search(r"jk=([a-zA-Z0-9]+)", href)
                        if m: jk = m.group(1)
                    # Method 4: parent div data-jk
                    if not jk:
                        try:
                            parent = driver.execute_script(
                                "return arguments[0].closest('[data-jk]')", a)
                            if parent:
                                jk = parent.get_attribute("data-jk") or ""
                        except Exception:
                            pass
                    if not jk: continue

                    link = f"https://{domain}/viewjob?jk={jk}"
                    # Only skip if already seen THIS run (not from CSV)
                    if link in seen_this_run: continue
                    seen_this_run.add(link)

                    title = ""
                    for sel in [f"span[id='jobTitle-{jk}']", "span[title]"]:
                        els = a.find_elements(By.CSS_SELECTOR, sel)
                        if els:
                            title = (els[0].get_attribute("title") or els[0].text).strip()
                            if title: break
                    if not title:
                        lbl = a.get_attribute("aria-label") or a.text or ""
                        title = lbl.split(" at ")[0].strip()
                    if not title: continue

                    # ── فیلتر عنوان: فقط شغل‌های مرتبط با keyword
                    if not title_matches_keyword(title, keywords[0] if keywords else ""):
                        continue

                    company, loc_text = "N/A", city or country
                    try:
                        card = driver.find_element(By.CSS_SELECTOR, f"div[data-jk='{jk}']")
                        for sel in ["[data-testid='company-name']", ".companyName"]:
                            el = card.find_elements(By.CSS_SELECTOR, sel)
                            if el: company = el[0].text.strip(); break
                        for sel in ["[data-testid='text-location']", ".companyLocation"]:
                            el = card.find_elements(By.CSS_SELECTOR, sel)
                            if el: loc_text = el[0].text.strip(); break
                    except Exception:
                        pass

                    jobs.append({"Title": title, "Company": company or "N/A",
                                 "Location": loc_text or city or country,
                                 "Source": "Indeed", "Link": link, "Date Found": now_str()})
                    page_new += 1
                except Exception as e:
                    log(f"[Indeed] card err: {e}")

            log(f"[Indeed] p{page+1}: +{page_new}")
            if page_new == 0 and page > 0: break
        except Exception as e:
            log(f"[Indeed] p{page+1} err: {e}"); break

    log(f"[Indeed] ✓ {len(jobs)} total")
    return jobs


def scrape_linkedin_api(keywords, country, city, existing, log, date_range=14, max_pages=10):
    """
    LinkedIn public jobs API — no login needed.
    10 pages × 25 = up to 250 jobs per run.
    """
    jobs = []
    # Build query: all keywords + Italian equivalents for broader coverage
    query    = " ".join(keywords)
    location = f"{city}, {country}" if city.strip() else country

    tpr_map = {7: "r604800", 14: "r1209600", 30: "r2592000", 90: "r7776000", 0: ""}
    tpr = tpr_map.get(date_range, "r1209600")

    headers = {
        "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                       "AppleWebKit/537.36 (KHTML, like Gecko) "
                       "Chrome/124.0.0.0 Safari/537.36"),
        "Accept-Language": "en-US,en;q=0.9,it;q=0.8",
        "Referer": "https://www.linkedin.com/jobs/",
    }

    seen_this_run = set()
    log(f"[LinkedIn] query='{query}'  location='{location}'  pages={max_pages}")
    for page in range(max_pages):
        start = page * 25
        params = {
            "keywords": query,
            "location": location,
            "start": start,
            "count": 25,
            "sortBy": "DD",
        }
        if tpr:
            params["f_TPR"] = tpr

        url = f"https://www.linkedin.com/jobs-guest/jobs/api/seeMoreJobPostings/search?{urlencode(params)}"
        log(f"[LinkedIn] API p{page+1}: offset={start}")

        try:
            resp = requests.get(url, headers=headers, timeout=20)
            if resp.status_code == 429:
                log("[LinkedIn] rate limited — waiting 15s")
                time.sleep(15)
                resp = requests.get(url, headers=headers, timeout=20)
            if not resp.ok:
                log(f"[LinkedIn] API error {resp.status_code}")
                break

            html = resp.text
            if not html.strip() or "<li>" not in html:
                log(f"[LinkedIn] empty page {page+1}")
                break

            # Parse job IDs from HTML snippet returned by API
            ids = re.findall(r'data-entity-urn="urn:li:jobPosting:(\d+)"', html)
            if not ids:
                ids = re.findall(r'"jobPostingId":(\d+)', html)
            if not ids:
                # fallback: extract from job links
                ids = re.findall(r'/jobs/view/(\d+)', html)

            ids = list(dict.fromkeys(ids))  # dedupe preserving order
            log(f"[LinkedIn] p{page+1}: {len(ids)} job IDs")

            if not ids:
                break

            # Parse job details from HTML using regex (more reliable than HTMLParser)
            # LinkedIn API returns HTML snippets like:
            # <h3 class="base-search-card__title">Title</h3>
            # <h4 class="base-search-card__subtitle"><a>Company</a></h4>
            # <span class="job-search-card__location">Location</span>

            # Extract all job cards as blocks
            card_blocks = re.findall(
                r'<li[^>]*>(.*?)</li>', html, re.DOTALL)

            parsed = {}
            for block in card_blocks:
                # Job ID
                jid_m = re.search(r'data-entity-urn="[^"]*:(\d+)"', block)
                if not jid_m:
                    jid_m = re.search(r'/jobs/view/(\d+)/', block)
                if not jid_m:
                    continue
                jid = jid_m.group(1)

                # Title — try multiple class names
                title = ""
                for pat in [
                    r'class="base-search-card__title"[^>]*>\s*(.*?)\s*</h3>',
                    r'class="[^"]*job[^"]*title[^"]*"[^>]*>\s*(.*?)\s*</',
                    r'aria-label="([^"]+)"',
                ]:
                    m = re.search(pat, block, re.DOTALL | re.IGNORECASE)
                    if m:
                        title = re.sub(r'<[^>]+>', '', m.group(1)).strip()
                        if title:
                            break

                # Company
                company = ""
                for pat in [
                    r'class="base-search-card__subtitle"[^>]*>.*?<a[^>]*>\s*(.*?)\s*</a>',
                    r'class="[^"]*company[^"]*"[^>]*>\s*(.*?)\s*</',
                ]:
                    m = re.search(pat, block, re.DOTALL | re.IGNORECASE)
                    if m:
                        company = re.sub(r'<[^>]+>', '', m.group(1)).strip()
                        if company:
                            break

                # Location
                loc = ""
                for pat in [
                    r'class="job-search-card__location"[^>]*>\s*(.*?)\s*</',
                    r'class="[^"]*location[^"]*"[^>]*>\s*(.*?)\s*</',
                ]:
                    m = re.search(pat, block, re.DOTALL | re.IGNORECASE)
                    if m:
                        loc = re.sub(r'<[^>]+>', '', m.group(1)).strip()
                        if loc:
                            break

                parsed[jid] = {
                    "title":    title or "",
                    "company":  company or "N/A",
                    "location": loc or location,
                }

            page_new = 0
            for jid in ids:
                link = f"https://www.linkedin.com/jobs/view/{jid}/"
                if link in seen_this_run:
                    continue
                seen_this_run.add(jid)

                info     = parsed.get(jid, {})
                title    = info.get("title", "").strip()
                company  = info.get("company", "N/A")
                loc_text = info.get("location", location)

                # Skip if we couldn't get a real title (just "Job 12345")
                if not title:
                    title = f"LinkedIn Job {jid}"

                # ── فیلتر عنوان
                if not title_matches_keyword(title, keywords[0] if keywords else ""):
                    continue

                jobs.append({"Title": title, "Company": company,
                             "Location": loc_text, "Source": "LinkedIn",
                             "Link": link, "Date Found": now_str()})
                existing.add(link)
                page_new += 1

            log(f"[LinkedIn] p{page+1}: +{page_new} (total {len(jobs)})")
            if page_new == 0 and page > 0:
                break
            time.sleep(1.5)  # polite rate limiting

        except Exception as e:
            log(f"[LinkedIn] p{page+1} err: {e}")
            break

    log(f"[LinkedIn] ✓ {len(jobs)} total")
    return jobs


def scrape_glassdoor(driver, keywords, country, city, existing, log, date_range=14):
    jobs = []
    query    = " ".join(keywords[:2])
    location = f"{city}, {country}" if city.strip() else country
    params = urlencode({"sc.keyword": query, "locKeyword": location, "locT": "N"})
    url = f"https://www.glassdoor.com/Job/jobs.htm?{params}"
    log(f"[Glassdoor] {url}")
    try:
        driver.get(url)
        time.sleep(5)
        for sel in ["button[data-test='modal-close-btn']",".modal_closeIcon-svg",
                    "button[alt='Close']","[class*='CloseButton']"]:
            try: driver.find_element(By.CSS_SELECTOR, sel).click(); time.sleep(1); break
            except Exception: pass
        scroll_down(driver, 3, 2)
        cards = driver.find_elements(By.CSS_SELECTOR,
            "li[data-jobid], div[data-jobid], article[data-jobid]")
        log(f"[Glassdoor] {len(cards)} cards")
        for card in cards[:40]:
            try:
                job_id = card.get_attribute("data-jobid") or ""
                if not job_id: continue
                link = f"https://www.glassdoor.com/job-listing/j?jl={job_id}"
                if link in existing: continue
                t_el = card.find_elements(By.CSS_SELECTOR,
                    "a[class*='jobTitle'], a[data-test='job-link'], span[class*='title']")
                title = t_el[0].text.strip() if t_el else ""
                if not title: continue
                c_el = card.find_elements(By.CSS_SELECTOR,
                    "span[class*='employerName'], div[class*='companyName']")
                company = c_el[0].text.strip() if c_el else "N/A"
                l_el = card.find_elements(By.CSS_SELECTOR,
                    "div[class*='location'], span[class*='location']")
                loc_text = l_el[0].text.strip() if l_el else location
                jobs.append({"Title": title, "Company": company,
                             "Location": loc_text, "Source": "Glassdoor",
                             "Link": link, "Date Found": now_str()})
                existing.add(link)
            except Exception as e:
                log(f"[Glassdoor] card err: {e}")
    except Exception as e:
        log(f"[Glassdoor] err: {e}")
    log(f"[Glassdoor] ✓ {len(jobs)}")
    return jobs


def scrape_infojobs(driver, keywords, country, city, existing, log):
    jobs = []
    query = "%20".join(keywords[:3])
    loc   = city or country
    url = f"https://www.infojobs.it/offerte-lavoro/offerte-lavoro.xhtml?keyword={quote_plus(' '.join(keywords[:3]))}&provinceOrCity={loc}&sortBy=PUBLICATION_DATE"
    log(f"[InfoJobs] {url}")
    try:
        driver.get(url); time.sleep(5); scroll_down(driver, 3, 2)
        cards = driver.find_elements(By.CSS_SELECTOR,
            "div[class*='sui-OfferCard'], div[class*='OfferCard'], article[class*='offer']")
        log(f"[InfoJobs] {len(cards)} cards")
        for card in cards[:40]:
            try:
                a_el = card.find_elements(By.CSS_SELECTOR, "a[href*='/offerta/']")
                if not a_el: continue
                link = a_el[0].get_attribute("href").split("?")[0]
                if link in existing: continue
                t_el = card.find_elements(By.CSS_SELECTOR, "h2, h3, [class*='title']")
                title = t_el[0].text.strip() if t_el else ""
                if not title: continue
                c_el = card.find_elements(By.CSS_SELECTOR, "[class*='company'], [class*='employer']")
                company = c_el[0].text.strip() if c_el else "N/A"
                l_el = card.find_elements(By.CSS_SELECTOR, "[class*='location'], [class*='city']")
                loc_text = l_el[0].text.strip() if l_el else loc
                jobs.append({"Title": title, "Company": company, "Location": loc_text,
                             "Source": "InfoJobs", "Link": link, "Date Found": now_str()})
                existing.add(link)
            except Exception as e:
                log(f"[InfoJobs] err: {e}")
    except Exception as e:
        log(f"[InfoJobs] err: {e}")
    log(f"[InfoJobs] ✓ {len(jobs)}")
    return jobs


def scrape_subito(driver, keywords, country, city, existing, log):
    jobs = []
    url = f"https://www.subito.it/annunci-italia/vendita/lavoro/?q={quote_plus(' '.join(keywords[:2]))}"
    log(f"[Subito] {url}")
    try:
        driver.get(url); time.sleep(5); scroll_down(driver, 2, 2)
        cards = driver.find_elements(By.CSS_SELECTOR,
            "div[class*='SmallCard'], article[class*='item']")
        log(f"[Subito] {len(cards)} cards")
        for card in cards[:30]:
            try:
                a_el = card.find_elements(By.CSS_SELECTOR, "a[href]")
                if not a_el: continue
                link = a_el[0].get_attribute("href").split("?")[0]
                if link in existing: continue
                t_el = card.find_elements(By.CSS_SELECTOR, "h2, h3, [class*='title']")
                title = t_el[0].text.strip() if t_el else ""
                if not title: continue
                l_el = card.find_elements(By.CSS_SELECTOR, "[class*='town'], [class*='location']")
                loc_text = l_el[0].text.strip() if l_el else city or "Italy"
                jobs.append({"Title": title, "Company": "N/A", "Location": loc_text,
                             "Source": "Subito.it", "Link": link, "Date Found": now_str()})
                existing.add(link)
            except Exception as e:
                log(f"[Subito] err: {e}")
    except Exception as e:
        log(f"[Subito] err: {e}")
    log(f"[Subito] ✓ {len(jobs)}")
    return jobs


def scrape_monster_it(driver, keywords, country, city, existing, log):
    jobs = []
    url = f"https://www.monster.it/lavoro/cerca/?q={quote_plus(' '.join(keywords[:3]))}&where={city or country}&sort=date.descend"
    log(f"[Monster IT] {url}")
    try:
        driver.get(url); time.sleep(5); scroll_down(driver, 3, 2)
        cards = driver.find_elements(By.CSS_SELECTOR,
            "div[data-jobid], article[class*='job-card'], div[class*='job-search-card']")
        log(f"[Monster IT] {len(cards)} cards")
        for card in cards[:40]:
            try:
                jid = card.get_attribute("data-jobid") or ""
                link = f"https://www.monster.it/lavoro/dettaglio/{jid}" if jid else ""
                if not link:
                    a_el = card.find_elements(By.CSS_SELECTOR, "a[href*='/lavoro/']")
                    if a_el: link = a_el[0].get_attribute("href").split("?")[0]
                if not link or link in existing: continue
                t_el = card.find_elements(By.CSS_SELECTOR, "h2, h3, [class*='title']")
                title = t_el[0].text.strip() if t_el else ""
                if not title: continue
                c_el = card.find_elements(By.CSS_SELECTOR, "[class*='company'], [class*='employer']")
                company = c_el[0].text.strip() if c_el else "N/A"
                l_el = card.find_elements(By.CSS_SELECTOR, "[class*='location']")
                loc_text = l_el[0].text.strip() if l_el else city or country
                jobs.append({"Title": title, "Company": company, "Location": loc_text,
                             "Source": "Monster IT", "Link": link, "Date Found": now_str()})
                existing.add(link)
            except Exception as e:
                log(f"[Monster IT] err: {e}")
    except Exception as e:
        log(f"[Monster IT] err: {e}")
    log(f"[Monster IT] ✓ {len(jobs)}")
    return jobs


# ══════════════════════════════════════════════════════════════════════════════
#  GUI  — Original 2-column layout with ttk.Notebook tabs
# ══════════════════════════════════════════════════════════════════════════════

BG      = "#F4F6F9"
CARD    = "#FFFFFF"
BORDER  = "#D1D9E6"
ACCENT  = "#2563EB"
ACCENT2 = "#1D4ED8"
TXT     = "#1E293B"
MUTED   = "#64748B"
GREEN   = "#16A34A"
RED     = "#DC2626"
AMBER   = "#D97706"
LOG_BG  = "#F8FAFC"
SUCCESS = GREEN
DANGER  = RED
WARNING = AMBER
TXT2    = MUTED
LOG_TXT = "#334155"


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Job Scraper Pro  v4")
        self.geometry("900x780")
        self.minsize(760, 620)
        self.configure(bg=BG)
        self._running = False
        self._cfg = load_config()
        self._build()

    # ─── helpers ──────────────────────────────────────────────────────────────

    def _card(self, p):
        return tk.Frame(p, bg=CARD,
                        highlightbackground=BORDER, highlightthickness=1)

    def _section(self, p, t):
        tk.Label(p, text=t.upper(), font=("Segoe UI", 7, "bold"),
                 bg=CARD, fg=MUTED).pack(anchor="w", padx=12, pady=(12, 2))
        tk.Frame(p, bg=BORDER, height=1).pack(fill="x", padx=12, pady=(0, 8))

    def _lbl(self, p, t):
        return tk.Label(p, text=t, font=("Segoe UI", 9), bg=CARD, fg=MUTED)

    def _entry(self, p, default=""):
        e = tk.Entry(p, bg=LOG_BG, fg=TXT, relief="flat",
                     font=("Segoe UI", 10),
                     highlightthickness=1,
                     highlightbackground=BORDER,
                     highlightcolor=ACCENT,
                     insertbackground=ACCENT)
        e.insert(0, default)
        return e

    # ─── logging ──────────────────────────────────────────────────────────────

    def log(self, msg: str):
        tag = "info"
        m = msg.lower()
        if any(x in m for x in ["✓", "✅", "done", "saved", "sent"]):
            tag = "ok"
        elif any(x in m for x in ["error", "fatal", "✗"]):
            tag = "err"
        elif any(x in m for x in ["⚠", "warn", "timeout", "rate"]):
            tag = "warn"
        elif "═" in msg or "─" in msg:
            tag = "muted"

        def _w():
            self.log_box.configure(state="normal")
            self.log_box.insert("end",
                f"{datetime.now().strftime('%H:%M:%S')}  {msg}\n", tag)
            self.log_box.see("end")
            self.log_box.configure(state="disabled")
        self.after(0, _w)

    def _set_status(self, text, color=GREEN):
        self.after(0, lambda: self.status_lbl.configure(
            text=f"⬤  {text}", fg=color))

    def _upd_counts(self, found, saved=None):
        self.after(0, lambda: self.count_lbl.configure(
            text=f"Found this run: {found}"))
        if saved is not None:
            self.after(0, lambda: self.saved_lbl.configure(
                text=f"New jobs saved: {saved}"))

    def _clear_log(self):
        self.log_box.configure(state="normal")
        self.log_box.delete("1.0", "end")
        self.log_box.configure(state="disabled")

    # ─── build ────────────────────────────────────────────────────────────────

    def _build(self):
        # Notebook style
        style = ttk.Style()
        style.theme_use("default")
        style.configure("TNotebook", background=BG, borderwidth=0)
        style.configure("TNotebook.Tab", background=BORDER, foreground=TXT,
                        font=("Segoe UI", 9), padding=[14, 6])
        style.map("TNotebook.Tab",
                  background=[("selected", CARD)],
                  foreground=[("selected", ACCENT)])
        style.configure("Jobs.Treeview",
                        background=CARD, fieldbackground=CARD,
                        foreground=TXT, font=("Segoe UI", 9), rowheight=26)
        style.configure("Jobs.Treeview.Heading",
                        background=BG, foreground=MUTED,
                        font=("Segoe UI", 9, "bold"))
        style.map("Jobs.Treeview", background=[("selected", ACCENT)])

        nb = ttk.Notebook(self)
        nb.pack(fill="both", expand=True, padx=14, pady=10)

        # ── Tab 1: Search
        t1 = tk.Frame(nb, bg=BG)
        nb.add(t1, text="  🔍 Search  ")
        self._build_search_tab(t1)

        # ── Tab 2: Notifications
        t2 = tk.Frame(nb, bg=BG)
        nb.add(t2, text="  🔔 Notifications  ")
        self._build_notif_tab(t2)

        # ── Tab 3: Results
        t3 = tk.Frame(nb, bg=BG)
        nb.add(t3, text="  📊 Results  ")
        self._build_results_tab(t3)

        # Switch to results after scrape
        self._nb = nb
        self._tab_results = t3

    # ─── TAB 1: Search ────────────────────────────────────────────────────────

    def _build_search_tab(self, parent):
        parent.columnconfigure(0, weight=3)
        parent.columnconfigure(1, weight=2)
        parent.rowconfigure(1, weight=1)

        # ── LEFT: form
        fc = self._card(parent)
        fc.grid(row=0, column=0, sticky="nsew", padx=(0, 8), pady=4)

        self._section(fc, "Keywords  (هر فیلد = یک جستجوی مستقل)")

        saved_kws = self._cfg.get("keywords_list",
            ["it support", "help desk", "desktop support", "supporto informatico"])
        if isinstance(saved_kws, str):
            # migrate from old comma format
            saved_kws = [k.strip() for k in saved_kws.split(",")]
        while len(saved_kws) < 4:
            saved_kws.append("")

        self.kw_entries = []
        kw_grid = tk.Frame(fc, bg=CARD)
        kw_grid.pack(fill="x", padx=12, pady=(0, 10))
        kw_grid.columnconfigure(1, weight=1)

        for i in range(4):
            color = ["#2563EB", "#16A34A", "#D97706", "#9333EA"][i]
            badge = tk.Label(kw_grid,
                text=f"  {i+1}  ",
                font=("Segoe UI", 9, "bold"),
                bg=color, fg="white")
            badge.grid(row=i, column=0, sticky="w", pady=3, padx=(0, 8))

            e = tk.Entry(kw_grid, bg=LOG_BG, fg=TXT, relief="flat",
                         font=("Segoe UI", 10),
                         highlightthickness=1,
                         highlightbackground=color,
                         highlightcolor=color,
                         insertbackground=color)
            e.insert(0, saved_kws[i] if i < len(saved_kws) else "")
            e.grid(row=i, column=1, sticky="ew", pady=3)
            self.kw_entries.append(e)

        self._section(fc, "Location & Date")
        geo = tk.Frame(fc, bg=CARD)
        geo.pack(fill="x", padx=12, pady=(0, 10))
        geo.columnconfigure(0, weight=1)
        geo.columnconfigure(1, weight=1)
        geo.columnconfigure(2, weight=1)

        lf = tk.Frame(geo, bg=CARD)
        lf.grid(row=0, column=0, sticky="ew", padx=(0, 5))
        self._lbl(lf, "Country").pack(anchor="w")
        self.country_entry = self._entry(lf, self._cfg.get("country", "Italy"))
        self.country_entry.pack(fill="x")

        rf = tk.Frame(geo, bg=CARD)
        rf.grid(row=0, column=1, sticky="ew", padx=(5, 5))
        self._lbl(rf, "City (optional)").pack(anchor="w")
        self.city_entry = self._entry(rf, self._cfg.get("city", "Roma"))
        self.city_entry.pack(fill="x")

        df = tk.Frame(geo, bg=CARD)
        df.grid(row=0, column=2, sticky="ew", padx=(5, 0))
        self._lbl(df, "Date Range").pack(anchor="w")
        self.date_var = tk.StringVar(value=self._cfg.get("date_range", "2 weeks"))
        ttk.Combobox(df, textvariable=self.date_var, state="readonly",
                     values=["1 week", "2 weeks", "1 month", "3 months", "Any time"],
                     font=("Segoe UI", 9)).pack(fill="x", ipady=3)

        self._section(fc, "Sites")
        sr = tk.Frame(fc, bg=CARD)
        sr.pack(fill="x", padx=12, pady=(0, 4))
        self.sites = {}
        all_sites = [
            ("Indeed",     "🌍 Indeed"),
            ("LinkedIn",   "💼 LinkedIn"),
            ("Glassdoor",  "🚪 Glassdoor"),
            ("InfoJobs",   "ℹ InfoJobs"),
            ("Subito.it",  "🛒 Subito"),
            ("Monster IT", "👹 Monster"),
        ]
        for i, (key, label) in enumerate(all_sites):
            var = tk.BooleanVar(value=self._cfg.get(f"site_{key}", True))
            tk.Checkbutton(sr, text=label, variable=var,
                           bg=CARD, fg=TXT, selectcolor=CARD,
                           activebackground=CARD, activeforeground=ACCENT,
                           font=("Segoe UI", 9), cursor="hand2"
                           ).grid(row=i//3, column=i%3, sticky="w",
                                  padx=(0, 16), pady=2)
            self.sites[key] = var

        self._section(fc, "Options")
        opt = tk.Frame(fc, bg=CARD)
        opt.pack(fill="x", padx=12, pady=(0, 14))

        self.headless_var = tk.BooleanVar(value=self._cfg.get("headless", True))
        tk.Checkbutton(opt, text="Headless browser",
                       variable=self.headless_var, bg=CARD, fg=MUTED,
                       selectcolor=CARD, activebackground=CARD,
                       font=("Segoe UI", 9), cursor="hand2"
                       ).pack(side="left", padx=(0, 18))

        self.docx_var = tk.BooleanVar(value=self._cfg.get("export_docx", True))
        lbl = "Save Word (.docx)" if DOCX_OK else "Word (pip install python-docx)"
        tk.Checkbutton(opt, text=lbl, variable=self.docx_var,
                       bg=CARD, fg=MUTED, selectcolor=CARD,
                       activebackground=CARD, font=("Segoe UI", 9),
                       cursor="hand2").pack(side="left")

        # ── RIGHT: status + actions
        rc = self._card(parent)
        rc.grid(row=0, column=1, sticky="nsew", pady=4)

        self._section(rc, "Status")
        self.status_lbl = tk.Label(rc, text="⬤  Ready",
                                   font=("Segoe UI", 11, "bold"), bg=CARD, fg=GREEN)
        self.status_lbl.pack(anchor="w", padx=12, pady=(0, 4))

        self.count_lbl = tk.Label(rc, text="Found this run: 0",
                                  font=("Segoe UI", 9), bg=CARD, fg=MUTED)
        self.count_lbl.pack(anchor="w", padx=12)

        self.saved_lbl = tk.Label(rc, text="New jobs saved: —",
                                  font=("Segoe UI", 9), bg=CARD, fg=MUTED)
        self.saved_lbl.pack(anchor="w", padx=12, pady=(2, 4))

        # Progress bar
        self.progress_var = tk.DoubleVar(value=0)
        prog = ttk.Progressbar(rc, variable=self.progress_var,
                               maximum=100, mode="determinate")
        prog.pack(fill="x", padx=12, pady=(2, 2))
        self.progress_lbl = tk.Label(rc, text="",
                                     font=("Segoe UI", 8), bg=CARD, fg=MUTED)
        self.progress_lbl.pack(anchor="w", padx=12, pady=(0, 10))

        self._section(rc, "Actions")

        self.run_btn = tk.Button(rc, text="▶  Start Scraping",
                                 command=self._start,
                                 bg=ACCENT, fg="white",
                                 font=("Segoe UI", 10, "bold"),
                                 relief="flat", padx=16, pady=8,
                                 cursor="hand2", activebackground=ACCENT2,
                                 activeforeground="white")
        self.run_btn.pack(fill="x", padx=12, pady=(0, 6))

        self.stop_btn = tk.Button(rc, text="■  Stop",
                                  command=self._stop,
                                  bg=BORDER, fg=MUTED,
                                  font=("Segoe UI", 10, "bold"),
                                  relief="flat", padx=16, pady=8,
                                  cursor="hand2", state="disabled")
        self.stop_btn.pack(fill="x", padx=12, pady=(0, 6))

        tk.Button(rc, text="📂  Open CSV",
                  command=self._open_csv,
                  bg=LOG_BG, fg=ACCENT, font=("Segoe UI", 9),
                  relief="flat", padx=16, pady=7, cursor="hand2"
                  ).pack(fill="x", padx=12, pady=(0, 3))

        tk.Button(rc, text="📄  Open Word Report",
                  command=self._open_docx,
                  bg=LOG_BG, fg=ACCENT, font=("Segoe UI", 9),
                  relief="flat", padx=16, pady=7, cursor="hand2"
                  ).pack(fill="x", padx=12, pady=(0, 14))

        # ── BOTTOM LOG (full width)
        lc = self._card(parent)
        lc.grid(row=1, column=0, columnspan=2, sticky="nsew", pady=(10, 0))

        log_top = tk.Frame(lc, bg=CARD)
        log_top.pack(fill="x", padx=12, pady=(10, 4))
        tk.Label(log_top, text="ACTIVITY LOG", font=("Segoe UI", 7, "bold"),
                 bg=CARD, fg=MUTED).pack(side="left")
        tk.Button(log_top, text="🗑 Clear", font=("Segoe UI", 8),
                  bg=LOG_BG, fg=MUTED, relief="flat", padx=6, pady=1,
                  cursor="hand2", command=self._clear_log).pack(side="right")
        tk.Frame(lc, bg=BORDER, height=1).pack(fill="x", padx=12, pady=(0, 6))

        self.log_box = scrolledtext.ScrolledText(
            lc, bg=LOG_BG, fg=LOG_TXT,
            font=("Consolas", 8), relief="flat",
            state="disabled", wrap="word")
        self.log_box.pack(fill="both", expand=True, padx=12, pady=(0, 12))
        self.log_box.tag_config("ok",    foreground=GREEN)
        self.log_box.tag_config("err",   foreground=RED)
        self.log_box.tag_config("warn",  foreground=AMBER)
        self.log_box.tag_config("info",  foreground=ACCENT)
        self.log_box.tag_config("muted", foreground=BORDER)

    # ─── TAB 2: Notifications ─────────────────────────────────────────────────

    def _build_notif_tab(self, parent):
        wrap = tk.Frame(parent, bg=BG)
        wrap.pack(fill="both", expand=True, padx=8, pady=8)

        # Telegram
        tg = self._card(wrap)
        tg.pack(fill="x", pady=(0, 10))
        self._section(tg, "Telegram Bot")

        tk.Label(tg,
            text="1. @BotFather → /newbot → copy Token\n"
                 "2. Start chat with your bot\n"
                 "3. Get Chat ID: send a message, then visit:\n"
                 "   https://api.telegram.org/bot<TOKEN>/getUpdates",
            font=("Segoe UI", 8), bg=CARD, fg=MUTED, justify="left"
        ).pack(anchor="w", padx=12, pady=(0, 8))

        tg_f = tk.Frame(tg, bg=CARD)
        tg_f.pack(fill="x", padx=12)
        tg_f.columnconfigure(1, weight=1)
        tk.Label(tg_f, text="Bot Token:", font=("Segoe UI", 9),
                 bg=CARD, fg=TXT, width=13, anchor="w"
                 ).grid(row=0, column=0, sticky="w", pady=3)
        self.tg_token = self._entry(tg_f, self._cfg.get("tg_token", ""))
        self.tg_token.grid(row=0, column=1, sticky="ew", padx=(8, 0))
        tk.Label(tg_f, text="Chat ID:", font=("Segoe UI", 9),
                 bg=CARD, fg=TXT, width=13, anchor="w"
                 ).grid(row=1, column=0, sticky="w", pady=3)
        self.tg_chat = self._entry(tg_f, self._cfg.get("tg_chat_id", ""))
        self.tg_chat.grid(row=1, column=1, sticky="ew", padx=(8, 0))

        tg_bot = tk.Frame(tg, bg=CARD)
        tg_bot.pack(fill="x", padx=12, pady=(8, 14))
        self.tg_en = tk.BooleanVar(value=self._cfg.get("tg_enabled", False))
        tk.Checkbutton(tg_bot, text="Send Telegram when scrape finishes",
                       variable=self.tg_en, bg=CARD, fg=TXT,
                       selectcolor=CARD, font=("Segoe UI", 9),
                       cursor="hand2").pack(side="left")
        tk.Button(tg_bot, text="Test Now", command=self._test_tg,
                  bg=LOG_BG, fg=ACCENT, font=("Segoe UI", 9),
                  relief="flat", padx=12, pady=4,
                  cursor="hand2").pack(side="right")

        # Email
        em = self._card(wrap)
        em.pack(fill="x", pady=(0, 10))
        self._section(em, "Email (SMTP)")

        tk.Label(em,
            text="Gmail App Password:\n"
                 "myaccount.google.com → Security → 2-Step Verification → App Passwords\n"
                 "Generate a 16-char password and paste below  |  Use Port: 465 (SSL)",
            font=("Segoe UI", 8), bg=CARD, fg=MUTED, justify="left"
        ).pack(anchor="w", padx=12, pady=(0, 8))

        em_f = tk.Frame(em, bg=CARD)
        em_f.pack(fill="x", padx=12)
        em_f.columnconfigure(1, weight=1)
        self.em_entries = {}
        fields = [
            ("SMTP Host",    "smtp_host",  "smtp.gmail.com"),
            ("SMTP Port",    "smtp_port",  "465"),
            ("Your Email",   "smtp_user",  ""),
            ("App Password", "smtp_pass",  ""),
            ("Send To",      "email_to",   ""),
        ]
        for i, (lbl, key, default) in enumerate(fields):
            tk.Label(em_f, text=lbl+":", font=("Segoe UI", 9),
                     bg=CARD, fg=TXT, width=13, anchor="w"
                     ).grid(row=i, column=0, sticky="w", pady=3)
            e = self._entry(em_f, self._cfg.get(key, default))
            if key == "smtp_pass":
                e.configure(show="●")
            e.grid(row=i, column=1, sticky="ew", padx=(8, 0))
            self.em_entries[key] = e

        em_bot = tk.Frame(em, bg=CARD)
        em_bot.pack(fill="x", padx=12, pady=(8, 14))
        self.em_en = tk.BooleanVar(value=self._cfg.get("email_enabled", False))
        tk.Checkbutton(em_bot, text="Send email when scrape finishes",
                       variable=self.em_en, bg=CARD, fg=TXT,
                       selectcolor=CARD, font=("Segoe UI", 9),
                       cursor="hand2").pack(side="left")
        tk.Button(em_bot, text="Test Now", command=self._test_email,
                  bg=LOG_BG, fg=ACCENT, font=("Segoe UI", 9),
                  relief="flat", padx=12, pady=4,
                  cursor="hand2").pack(side="right")

        # Save button
        tk.Button(wrap, text="💾  Save All Settings",
                  command=self._save_settings,
                  bg=ACCENT, fg="white",
                  font=("Segoe UI", 10, "bold"),
                  relief="flat", padx=20, pady=8,
                  cursor="hand2").pack(pady=6)

    # ─── TAB 3: Results ───────────────────────────────────────────────────────

    def _build_results_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)

        # ── Filter bar (top)
        fbar = self._card(parent)
        fbar.grid(row=0, column=0, sticky="ew", padx=4, pady=(4, 6))

        top_row = tk.Frame(fbar, bg=CARD)
        top_row.pack(fill="x", padx=12, pady=(10, 6))
        top_row.columnconfigure(1, weight=1)

        tk.Label(top_row, text="🔍 Search:", font=("Segoe UI", 9),
                 bg=CARD, fg=MUTED).grid(row=0, column=0, sticky="w")
        self.filter_entry = self._entry(top_row, "")
        self.filter_entry.grid(row=0, column=1, sticky="ew", padx=(8, 0))
        self.filter_entry.bind("<KeyRelease>", lambda e: self._refresh_results())

        # Source radio filter
        src_row = tk.Frame(fbar, bg=CARD)
        src_row.pack(fill="x", padx=12, pady=(0, 6))
        tk.Label(src_row, text="Source:", font=("Segoe UI", 8),
                 bg=CARD, fg=MUTED).pack(side="left")
        self.filter_source = tk.StringVar(value="All")
        for src in ["All", "Indeed", "LinkedIn", "Glassdoor",
                    "InfoJobs", "Subito.it", "Monster IT"]:
            tk.Radiobutton(src_row, text=src, variable=self.filter_source,
                           value=src, bg=CARD, fg=TXT, selectcolor=CARD,
                           font=("Segoe UI", 8), cursor="hand2",
                           command=self._refresh_results
                           ).pack(side="left", padx=(6, 0))

        # Action buttons
        btn_row = tk.Frame(fbar, bg=CARD)
        btn_row.pack(fill="x", padx=12, pady=(0, 10))
        tk.Button(btn_row, text="↺ Refresh",
                  command=lambda: [
                      self.filter_source.set("All"),
                      self._refresh_results()
                  ],
                  bg=LOG_BG, fg=ACCENT, font=("Segoe UI", 8),
                  relief="flat", padx=10, pady=4, cursor="hand2"
                  ).pack(side="left", padx=(0, 6))
        tk.Button(btn_row, text="🕐 Last Run Only",
                  command=lambda: [
                      self.filter_source.set("All"),
                      self._refresh_results(last_run_only=True)
                  ],
                  bg="#EFF6FF", fg=ACCENT, font=("Segoe UI", 8, "bold"),
                  relief="flat", padx=10, pady=4, cursor="hand2"
                  ).pack(side="left", padx=(0, 6))
        tk.Button(btn_row, text="🗑 Clear All Data",
                  command=self._clear_all_data,
                  bg="#FEE2E2", fg=RED, font=("Segoe UI", 8, "bold"),
                  relief="flat", padx=10, pady=4, cursor="hand2"
                  ).pack(side="left", padx=(0, 6))
        tk.Button(btn_row, text="📊 Open Excel", command=self._open_xlsx,
                  bg=LOG_BG, fg=ACCENT, font=("Segoe UI", 8),
                  relief="flat", padx=10, pady=4, cursor="hand2"
                  ).pack(side="left", padx=(0, 6))
        tk.Button(btn_row, text="📄 Open Word", command=self._open_docx,
                  bg=LOG_BG, fg=ACCENT, font=("Segoe UI", 8),
                  relief="flat", padx=10, pady=4, cursor="hand2"
                  ).pack(side="left")
        self.results_lbl = tk.Label(btn_row, text="",
                                    font=("Segoe UI", 8), bg=CARD, fg=MUTED)
        self.results_lbl.pack(side="right")

        # ── Treeview
        tree_card = self._card(parent)
        tree_card.grid(row=1, column=0, sticky="nsew", padx=4, pady=(0, 4))

        cols = ("Title", "Company", "Location", "Source", "Time", "Link")
        self.tree = ttk.Treeview(tree_card, columns=cols,
                                 show="headings", style="Jobs.Treeview")
        col_w = [230, 150, 110, 80, 130, 0]   # Link width=0 (hidden, stored as tag)
        for col, w in zip(cols, col_w):
            self.tree.heading(col, text=col,
                              command=lambda c=col: self._sort_tree(c))
            self.tree.column(col, width=w, minwidth=40,
                             stretch=(col != "Link"))

        vsb = ttk.Scrollbar(tree_card, orient="vertical",
                            command=self.tree.yview)
        hsb = ttk.Scrollbar(tree_card, orient="horizontal",
                            command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set,
                            xscrollcommand=hsb.set)
        hsb.pack(side="bottom", fill="x")
        vsb.pack(side="right",  fill="y")
        self.tree.pack(side="left", fill="both", expand=True)

        self.tree.bind("<Double-1>", self._open_job_link)
        self.tree.bind("<Return>",   self._open_job_link)

        # Tooltip for link
        self._tooltip = tk.Label(self, bg="#FFFDE7", fg=TXT,
                                 font=("Segoe UI", 8), relief="solid",
                                 borderwidth=1, wraplength=500)
        self.tree.bind("<Motion>", self._show_tooltip)
        self.tree.bind("<Leave>",  lambda e: self._tooltip.place_forget())

    # ─── results helpers ──────────────────────────────────────────────────────

    def _refresh_results(self, last_run_only=False):
        if not hasattr(self, "tree") or not hasattr(self, "filter_entry"):
            return
        if not os.path.exists(CSV_FILE):
            if hasattr(self, "results_lbl"):
                self.results_lbl.configure(text=f"No CSV found at: {CSV_FILE}")
            return
        try:
            from datetime import datetime as dt

            df = pd.read_csv(CSV_FILE).fillna("")
            total_all = len(df)

            # Normalize ALL date formats to a single sortable format
            def normalize_date(val):
                val = str(val).strip()
                for fmt in ["%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M",
                            "%m/%d/%Y %H:%M", "%m/%d/%Y %H:%M:%S",
                            "%d/%m/%Y %H:%M", "%d-%m-%Y %H:%M"]:
                    try:
                        return dt.strptime(val[:16], fmt[:len(fmt)])
                    except Exception:
                        pass
                return dt.min

            if "Date Found" in df.columns:
                df["_dt"] = df["Date Found"].apply(normalize_date)
                df = df.sort_values("_dt", ascending=False).reset_index(drop=True)

            # Last run only
            if last_run_only and len(df) > 0:
                from datetime import timedelta
                # Use Run column if available (most reliable)
                if "Run" in df.columns and df["Run"].nunique() > 1:
                    latest_run = df["Run"].iloc[0]
                    df = df[df["Run"] == latest_run].reset_index(drop=True)
                else:
                    # Fallback: find biggest gap
                    dts = df["_dt"].tolist()
                    if len(dts) > 1:
                        gaps = [(dts[i] - dts[i+1], i) for i in range(len(dts)-1)]
                        max_gap, max_idx = max(gaps, key=lambda x: x[0])
                        if max_gap > timedelta(minutes=5):
                            df = df.iloc[:max_idx+1].reset_index(drop=True)
                self.log(f"[Results] Last run: {len(df)} jobs")

            # text filter
            q = self.filter_entry.get().lower().strip()
            if q:
                df = df[df.apply(lambda r: q in str(r).lower(), axis=1)]

            # source filter
            src = self.filter_source.get() if hasattr(self, "filter_source") else "All"
            if src != "All":
                df = df[df["Source"].astype(str) == src]

            self.tree.delete(*self.tree.get_children())

            last_hour = None
            for _, row in df.iterrows():
                d = row.get("_dt", dt.min)
                # Separator when hour changes (= new run)
                hour_key = d.strftime("%Y-%m-%d %H") if d != dt.min else ""
                if hour_key != last_hour:
                    if last_hour is not None:
                        self.tree.insert("", "end",
                            values=("─"*25, "", "", "", "──────────", ""),
                            tags=("separator",))
                    last_hour = hour_key

                # Always show date as DD/MM/YYYY HH:MM
                if d != dt.min:
                    time_display = d.strftime("%d/%m/%Y  %H:%M")
                else:
                    time_display = _safe(row.get("Date Found", ""))[:16]

                link = _safe(row.get("Link", ""))
                self.tree.insert("", "end",
                    values=(
                        _safe(row.get("Title")),
                        _safe(row.get("Company")),
                        _safe(row.get("Location")),
                        _safe(row.get("Source")),
                        time_display,
                        link,
                    ),
                    tags=(link,)
                )

            self.tree.tag_configure("separator",
                background="#E2E8F0", foreground=MUTED)

            shown = len(df)
            extra = f"  (of {total_all} total)" if shown < total_all else ""
            if hasattr(self, "results_lbl"):
                self.results_lbl.configure(text=f"✓ {shown} job(s){extra}")
        except Exception as e:
            if hasattr(self, "results_lbl"):
                self.results_lbl.configure(text=f"Error: {e}")
            self.log(f"[Results] refresh error: {e}")

    def _sort_tree(self, col):
        rows = [(self.tree.set(k, col), k)
                for k in self.tree.get_children("")]
        rows.sort()
        for i, (_, k) in enumerate(rows):
            self.tree.move(k, "", i)

    def _open_job_link(self, event=None):
        sel = self.tree.selection()
        if not sel:
            return
        tags = self.tree.item(sel[0], "tags")
        link = tags[0] if tags else ""
        if link.startswith("http"):
            import webbrowser
            webbrowser.open(link)

    def _show_tooltip(self, event):
        item = self.tree.identify_row(event.y)
        if not item:
            self._tooltip.place_forget()
            return
        tags = self.tree.item(item, "tags")
        link = tags[0] if tags else ""
        if link.startswith("http"):
            self._tooltip.configure(text=f"🔗 {link}")
            x = self.tree.winfo_rootx() - self.winfo_rootx() + 20
            y = event.y + 28
            self._tooltip.place(x=x, y=y)
        else:
            self._tooltip.place_forget()

    # ─── file openers ─────────────────────────────────────────────────────────

    def _clear_all_data(self):
        ans = messagebox.askyesno(
            "Clear All Data",
            "این کار همه نتایج ذخیره‌شده رو پاک میکنه.\n"
            "(CSV, Excel, Word)\n\n"
            "مطمئنی؟")
        if not ans:
            return
        for f in [CSV_FILE, XLSX_FILE, DOCX_FILE]:
            try:
                if os.path.exists(f):
                    os.remove(f)
            except Exception as e:
                self.log(f"[Clear] could not delete {f}: {e}")
        self.tree.delete(*self.tree.get_children())
        if hasattr(self, "results_lbl"):
            self.results_lbl.configure(text="All data cleared ✓")
        self.log("🗑 All data cleared — ready for fresh scrape")

    def _open_xlsx(self):
        if os.path.exists(XLSX_FILE):
            os.startfile(XLSX_FILE) if os.name == "nt" else os.system(f'xdg-open "{XLSX_FILE}"')
        else:
            messagebox.showinfo("Info", f"No Excel file yet.\n\nExpected location:\n{XLSX_FILE}")

    def _open_csv(self):
        if os.path.exists(CSV_FILE):
            os.startfile(CSV_FILE) if os.name == "nt" else os.system(f'xdg-open "{CSV_FILE}"')
        else:
            messagebox.showinfo("Info", f"No CSV yet.\n\nExpected location:\n{CSV_FILE}")

    def _open_docx(self):
        if os.path.exists(DOCX_FILE):
            os.startfile(DOCX_FILE) if os.name == "nt" else os.system(f'xdg-open "{DOCX_FILE}"')
        else:
            messagebox.showinfo("Info", f"No Word report yet.\n\nExpected location:\n{DOCX_FILE}")

    # ─── settings ─────────────────────────────────────────────────────────────

    def _save_settings(self):
        cfg = {
            "keywords_list":  [e.get().strip() for e in self.kw_entries],
            "country":       self.country_entry.get().strip(),
            "city":          self.city_entry.get().strip(),
            "date_range":    self.date_var.get(),
            "headless":      self.headless_var.get(),
            "export_docx":   self.docx_var.get(),
            "tg_enabled":    self.tg_en.get(),
            "tg_token":      self.tg_token.get().strip(),
            "tg_chat_id":    self.tg_chat.get().strip(),
            "email_enabled": self.em_en.get(),
            **{k: e.get().strip() for k, e in self.em_entries.items()},
        }
        for site, var in self.sites.items():
            cfg[f"site_{site}"] = var.get()
        save_config(cfg)
        self._cfg = cfg
        messagebox.showinfo("Saved", "Settings saved ✓")

    def _test_tg(self):
        self._save_settings()
        ok = send_telegram(
            [{"Title": "Test", "Company": "Co", "Location": "Rome",
              "Source": "Test", "Link": "https://example.com"}],
            self._cfg, self.log)
        messagebox.showinfo("Telegram", "✓ Sent!" if ok else "✗ Failed — check log")

    def _test_email(self):
        self._save_settings()
        ok = send_email(
            [{"Title": "Test", "Company": "Co", "Location": "Rome",
              "Source": "Test", "Link": "https://example.com"}],
            self._cfg, self.log)
        messagebox.showinfo("Email", "✓ Sent!" if ok else "✗ Failed — check log")

    # ─── scrape control ───────────────────────────────────────────────────────

    def _start(self):
        # Collect non-empty keywords from 4 fields
        kws = [e.get().strip() for e in self.kw_entries if e.get().strip()]
        if not kws:
            messagebox.showerror("Error", "Enter at least one keyword"); return
        sites = [s for s, v in self.sites.items() if v.get()]
        if not sites:
            messagebox.showerror("Error", "Select at least one site"); return

        self._running = True
        self.run_btn.configure(state="disabled")
        self.stop_btn.configure(state="normal", bg="#FEE2E2", fg=RED)
        self._set_status("Running…", ACCENT)
        self._upd_counts(0)
        self.after(0, lambda: self.progress_var.set(0))
        self.after(0, lambda: self.progress_lbl.configure(text=""))

        threading.Thread(target=self._worker, args=(kws, sites), daemon=True).start()

    def _reset_ui(self, status_text="Ready", status_color=None):
        """Reset buttons to initial state — always called at end of run."""
        color = status_color or GREEN
        self.after(0, lambda: self.run_btn.configure(state="normal"))
        self.after(0, lambda: self.stop_btn.configure(
            state="disabled", bg=BORDER, fg=MUTED))
        self._set_status(status_text, color)

    def _stop(self):
        self._running = False
        self.log("⚠ Stop requested — finishing current step…")
        self._set_status("Stopping…", AMBER)
        # Reset buttons immediately so user can start again
        self.after(1500, lambda: self._reset_ui("Stopped", AMBER)
                   if not self._running else None)

    # ─── worker ───────────────────────────────────────────────────────────────

    def _worker(self, keywords, selected_sites):
        country    = self.country_entry.get().strip() or "Italy"
        city       = self.city_entry.get().strip()
        headless   = self.headless_var.get()
        dr_map     = {"1 week": 7, "2 weeks": 14, "1 month": 30,
                      "3 months": 90, "Any time": 0}
        date_range = dr_map.get(self.date_var.get(), 14)
        total_sites = len(selected_sites)

        existing = load_existing_links()
        all_jobs = []

        def upd_progress(done, label=""):
            pct = (done / total_sites) * 100
            self.after(0, lambda: self.progress_var.set(pct))
            self.after(0, lambda: self.progress_lbl.configure(text=label))

        self.log(f"{'═'*50}")
        self.log(f"  🚀 Started  {now_str()}")
        self.log(f"  Keywords : {' | '.join(keywords)}")
        self.log(f"  Location : {city + ', ' if city else ''}{country}")
        self.log(f"  Range    : {self.date_var.get()}")
        self.log(f"  Sites    : {', '.join(selected_sites)}")
        self.log(f"{'═'*50}")

        needs_browser = [s for s in selected_sites if s != "LinkedIn"]
        driver = None

        try:
            # ── LinkedIn — جداگانه برای هر keyword
            if "LinkedIn" in selected_sites and self._running:
                idx = selected_sites.index("LinkedIn") + 1
                upd_progress(idx - 0.5, f"LinkedIn [{idx}/{total_sites}]")
                self.log(f"\n▶ LinkedIn  [{idx}/{total_sites}]")
                ln_total = 0
                for kw in keywords:
                    if not self._running: break
                    self.log(f"  🔑 \"{kw}\"")
                    jobs = scrape_linkedin_api([kw], country, city, existing,
                                               self.log, date_range=date_range)
                    all_jobs.extend(jobs)
                    ln_total += len(jobs)
                    self._upd_counts(len(all_jobs))
                    self.log(f"     └─ {len(jobs)} jobs")
                upd_progress(idx, f"LinkedIn ✓ {ln_total}")

            # ── Selenium sites — جداگانه برای هر keyword
            if needs_browser and self._running:
                self.log("\n  Starting Chrome…")
                driver = make_driver(headless)

                fn_map = {
                    "Indeed":     lambda kw: scrape_indeed(driver, [kw], country, city, existing, self.log, date_range=date_range),
                    "Glassdoor":  lambda kw: scrape_glassdoor(driver, [kw], country, city, existing, self.log, date_range=date_range),
                    "InfoJobs":   lambda kw: scrape_infojobs(driver, [kw], country, city, existing, self.log),
                    "Subito.it":  lambda kw: scrape_subito(driver, [kw], country, city, existing, self.log),
                    "Monster IT": lambda kw: scrape_monster_it(driver, [kw], country, city, existing, self.log),
                }

                for site in needs_browser:
                    if not self._running: break
                    idx = selected_sites.index(site) + 1
                    upd_progress(idx - 0.5, f"{site} [{idx}/{total_sites}]")
                    self.log(f"\n▶ {site}  [{idx}/{total_sites}]")
                    self._set_status(f"{site}…", ACCENT)
                    site_total = 0
                    for kw in keywords:
                        if not self._running: break
                        self.log(f"  🔑 \"{kw}\"")
                        jobs = fn_map[site](kw)
                        all_jobs.extend(jobs)
                        site_total += len(jobs)
                        self._upd_counts(len(all_jobs))
                        self.log(f"     └─ {len(jobs)} jobs  (total: {len(all_jobs)})")
                    upd_progress(idx, f"{site} ✓ {site_total}")

        except Exception as e:
            self.log(f"FATAL: {e}")
        finally:
            if driver:
                try: driver.quit()
                except Exception: pass
                self.log("\n  Browser closed.")

        # Save — هر run جداگانه ذخیره میشه
        run_id = now_str()[:16]  # "2026-04-26 20:41"
        new_count = save_to_csv(all_jobs, run_id=run_id)
        self.log(f"\n{'═'*50}")
        self.log(f"  ✅ DONE  {now_str()}")
        self.log(f"  Scraped : {len(all_jobs)} total  |  New: {new_count}")
        self.log(f"  CSV     : {CSV_FILE}")
        self.log(f"{'═'*50}")

        # Word
        if self.docx_var.get():
            try:
                src = pd.read_csv(CSV_FILE) if os.path.exists(CSV_FILE) \
                      else pd.DataFrame(all_jobs)
                save_to_docx(src.to_dict("records"), self.log)
                self.log(f"  Word    : {DOCX_FILE}")
            except Exception as e:
                self.log(f"  [DOCX] ✗ {e}")

        # Excel
        try:
            src = pd.read_csv(CSV_FILE) if os.path.exists(CSV_FILE) \
                  else pd.DataFrame(all_jobs)
            save_to_xlsx(src.to_dict("records"), self.log)
            self.log(f"  Excel   : {XLSX_FILE}")
        except Exception as e:
            self.log(f"  [Excel] ✗ {e}")

        # Notify
        if all_jobs:
            if self._cfg.get("tg_enabled"):
                send_telegram(all_jobs, self._cfg, self.log)
            if self._cfg.get("email_enabled"):
                send_email(all_jobs, self._cfg, self.log)

        # UI reset — always runs whether stopped or finished normally
        self._running = False
        if all_jobs or new_count >= 0:
            self.after(0, lambda: self.progress_var.set(100))
            self.after(0, lambda: self.progress_lbl.configure(
                text=f"Done ✓  {new_count} new saved"))
            self._upd_counts(len(all_jobs), new_count)
            self._reset_ui(f"Done  ({new_count} new)", GREEN)
        else:
            self._reset_ui("Stopped", AMBER)
        self.after(0, self._refresh_results)


# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    app = App()
    app.mainloop()